#!/usr/bin/env python3
"""
PDF Text Extraction for Job Descriptions

Extracts text content from JD files using multiple methods:
- PyMuPDF (fitz) for standard PDFs  
- python-docx for Word documents
- OCR for scanned/image PDFs
"""

import os
import time
import pymupdf as fitz  # PyMuPDF
import docx
from typing import Dict, Any

def extract_pdf_text(file_path: str) -> Dict[str, Any]:
    """Extract text from PDF file using PyMuPDF"""
    try:
        start_time = time.time()
        
        doc = fitz.open(file_path)
        all_text = []
        page_count = len(doc)
        
        for page_index in range(page_count):
            page = doc[page_index]
            page_text = page.get_text("text")
            if page_text.strip():  # Only add non-empty pages
                all_text.append(f"\n\n--- Page {page_index + 1} ---\n{page_text.strip()}")

        doc.close()
        final_text = "\n".join(all_text).strip()
        
        elapsed = time.time() - start_time
        
        # Return success even if text is short - let caller decide if it's acceptable
        return {
            'success': True,
            'text': final_text,
            'extraction_method': 'pymupdf',
            'processing_time': elapsed,
            'pages': page_count,
            'text_length': len(final_text)
        }
            
    except Exception as e:
        return {
            'success': False,
            'text': '',
            'error': f"PDF extraction failed: {str(e)}"
        }

def extract_docx_text(file_path: str) -> Dict[str, Any]:
    """Extract text from DOCX file using python-docx"""
    try:
        start_time = time.time()
        
        doc = docx.Document(file_path)
        text_content = '\n'.join([para.text for para in doc.paragraphs])
        
        elapsed = time.time() - start_time
        
        if text_content.strip():
            return {
                'success': True,
                'text': text_content,
                'extraction_method': 'python-docx',
                'processing_time': elapsed,
                'paragraphs': len(doc.paragraphs)
            }
        else:
            return {
                'success': False,
                'text': '',
                'error': 'No text extracted from DOCX'
            }
            
    except Exception as e:
        return {
            'success': False,
            'text': '',
            'error': f"DOCX extraction failed: {str(e)}"
        }

def extract_with_ocr(file_path: str) -> Dict[str, Any]:
    """Extract text using OCR for image-based files"""
    try:
        import pytesseract
        from PIL import Image
        
        start_time = time.time()
        
        # For PDF files, convert to image first
        if file_path.lower().endswith('.pdf'):
            doc = fitz.open(file_path)
            all_text = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                pix = page.get_pixmap()
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # OCR the image
                page_text = pytesseract.image_to_string(img)
                all_text.append(f"\n\n--- Page {page_num + 1} (OCR) ---\n{page_text.strip()}")
            
            doc.close()
            final_text = "\n".join(all_text).strip()
            
        else:
            # Direct image OCR
            img = Image.open(file_path)
            final_text = pytesseract.image_to_string(img)
        
        elapsed = time.time() - start_time
        
        if final_text.strip():
            return {
                'success': True,
                'text': final_text,
                'extraction_method': 'ocr',
                'processing_time': elapsed
            }
        else:
            return {
                'success': False,
                'text': '',
                'error': 'No text extracted via OCR'
            }
            
    except ImportError:
        return {
            'success': False,
            'text': '',
            'error': 'OCR dependencies not installed (pytesseract, Pillow)'
        }
    except Exception as e:
        return {
            'success': False,
            'text': '',
            'error': f"OCR extraction failed: {str(e)}"
        }

def get_jd_pdf_path(jd_pdf_filename: str, base_path: str = '/app') -> str:
    """Get full path to JD PDF file"""
    return os.path.join(base_path, 'uploads', 'jds', jd_pdf_filename)


def extract_combined_text(job_data: dict, base_path: str = '/app') -> Dict[str, Any]:
    """
    Extract and combine text from PDF and JD text fields
    
    Args:
        job_data: Job data dict with jd_pdf_filename and jd_text fields
        base_path: Base uploads directory
        
    Returns: {
        'success': bool,
        'text': str,
        'sources': list,  # ['pdf', 'jd_text']
        'char_count': int,
        'error': str or None
    }
    """
    combined_text = ""
    sources = []
    
    # Extract from PDF if filename exists
    jd_pdf_filename = job_data.get('jd_pdf_filename', '')
    if jd_pdf_filename:
        file_path = get_jd_pdf_path(jd_pdf_filename, base_path)
        
        if os.path.exists(file_path):
            text_result = process_jd_file(file_path)
            
            if text_result.get('success'):
                pdf_text = text_result.get('text', '')
                if pdf_text:
                    combined_text += pdf_text
                    sources.append('pdf')
    
    # Add jd_text if exists
    jd_text = job_data.get('jd_text', '')
    if jd_text:
        if combined_text:
            combined_text += "\n\n" + jd_text
        else:
            combined_text = jd_text
        sources.append('jd_text')
    
    if not combined_text:
        return {
            'success': False,
            'text': '',
            'sources': [],
            'char_count': 0,
            'error': 'No JD text or PDF content found'
        }
    
    return {
        'success': True,
        'text': combined_text,
        'sources': sources,
        'char_count': len(combined_text),
        'error': None
    }


def process_jd_file(file_path: str) -> Dict[str, Any]:
    """
    Main function to extract text from JD file
    Returns: {
        'success': bool,
        'text': str,
        'extraction_method': str,
        'error': str or None
    }
    """
    if not os.path.exists(file_path):
        return {
            'success': False,
            'text': '',
            'error': f'File not found: {file_path}'
        }
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    # Try extraction based on file type
    if file_ext == '.pdf':
        # Try PyMuPDF first
        result = extract_pdf_text(file_path)
        
        # Only try OCR if PyMuPDF completely fails (not just poor quality)
        if not result['success']:
            #print(f"⚠️ PyMuPDF extraction failed, trying OCR...")
            ocr_result = extract_with_ocr(file_path)
            if ocr_result['success']:
                return ocr_result
            else:
                # If OCR also fails, return the original PyMuPDF error
                return result
        

        return result
        
    elif file_ext in ['.docx', '.doc']:
        return extract_docx_text(file_path)
        
    elif file_ext == '.txt':
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return {
                'success': True,
                'text': text,
                'extraction_method': 'plain_text'
            }
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'error': f"Text file reading failed: {str(e)}"
            }
    
    else:
        return {
            'success': False,
            'text': '',
            'error': f'Unsupported file format: {file_ext}'
        }